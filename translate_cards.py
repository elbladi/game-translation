from PIL import Image
import pytesseract
from googletrans import Translator
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
from datetime import datetime

result = []
translations = []
data = {}

def get_text_from_image(image_path):
    img = Image.open(image_path)
    text = pytesseract.image_to_string(img)
    text = text.replace('\n', ' ').replace('_', '').replace('’', '').replace('‘', '')
    return text

def create_card(doc, text):
    # Add a table with one cell (1 row, 1 column)
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)

    # Set the width and height of the cell
    cell.width = Pt(6 * 28.35)  
    cell.height = Pt(9 * 28.35) 

    # Add some text to the cell
    cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.text = text 

    font = cell.paragraphs[0].runs[0].font
    font.color.rgb = RGBColor(79, 113, 190)
    font.size = Pt(18)

    # Set a cell background (shading) color to RGB 4f71be.
    # shading_elm = parse_xml(r'<w:shd {} w:fill="4f71be"/>'.format(nsdecls('w')))
    # cell._tc.get_or_add_tcPr().append(shading_elm)

    doc.add_paragraph()

def create_doc_file():
    doc = Document()
    # Create a card for each element in the array
    for text in translations:
        create_card(doc, text)

    # Save the Word document
    doc.save("output.docx")

def translate_text():
    translator = Translator()
    for index, value in enumerate(result):
        translation = translator.translate(value, dest='es')
        print(f'Translation --> {translation.text}')
        data[index + 1] = {
            "english": value,
            "spanish": translation.text
        }
        translations.append(translation.text)

# folder_name = "tests"
folder_name = "images"

def read_images():
    # get list of images names
    files = os.listdir(f'./{folder_name}')
    for file_name in files:
        if(file_name.endswith(".png") or file_name.endswith(".jpeg") or file_name.endswith(".jpg")):
            text = get_text_from_image(f'./{folder_name}/{file_name}')
            print(text)
            result.append(text)

def create_data_file():
    with open('data.json', 'w') as json_file:
        json.dump(data, json_file, indent=2)

def print_time(execOn):
    current_time = datetime.now()
    human_readable_time = current_time.strftime("%Y-%m-%d %H:%M:%S")
    print(f"Execution {execOn} at:", human_readable_time)

try:
    print_time("started")
    read_images()
    translate_text()
    create_doc_file()
    create_data_file()
    print_time("end")
except:
    print("el horrorr")
    print_time("end with error")