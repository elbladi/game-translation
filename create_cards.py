from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import json

def create_card(doc, text):
    # Add a table with one cell (1 row, 1 column)
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)

    # Set the width and height of the cell
    cell.width = Pt(6 * 28.35)  
    cell.height = Pt(9 * 28.35) 

    # Add some text to the cell
    cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.text = text.strip()

    font = cell.paragraphs[0].runs[0].font
    font.color.rgb = RGBColor(79, 113, 190)
    font.size = Pt(18)

    # Set a cell background (shading) color to RGB 4f71be.
    # shading_elm = parse_xml(r'<w:shd {} w:fill="4f71be"/>'.format(nsdecls('w')))
    # cell._tc.get_or_add_tcPr().append(shading_elm)

    doc.add_paragraph()

translations = [] 

def create_doc_file():
    doc = Document()
    # Create a card for each element in the array
    for text in translations:
        create_card(doc, text)

    # Save the Word document
    doc.save("cards.docx")

def get_data():
    file_name = "data.json"
    if not os.path.exists(file_name):
        # I should automate this, maybe.
        raise ValueError(f"{file_name} not found in root path. Make sure to run generate_data.py first.")
    
    with open(file_name, "r") as file:
        data = json.load(file)

    for _, value in data.items():
        translations.append(value["spanish"])

get_data()
create_doc_file()